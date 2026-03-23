from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(run, font_name_en='Times New Roman', font_name_cn='宋体', size=None, bold=False):
    run.font.name = font_name_en
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
    if size:
        run.font.size = size
    if bold:
        run.font.bold = True

def update_style(doc, style_name, font_name_cn, font_size_pt, align=WD_ALIGN_PARAGRAPH.LEFT, 
                 line_spacing=1.25, space_before_pt=0, space_after_pt=0, 
                 indent_chars=0, bold=False, heading_level=None):
    
    # Get or create style
    try:
        if heading_level:
            style = doc.styles[style_name] # Built-in styles like 'Heading 1'
        else:
            style = doc.styles.add_style(style_name, 1) # 1 = Paragraph
    except:
        try:
            style = doc.styles[style_name]
        except:
            style = doc.styles.add_style(style_name, 1)

    # Font settings
    font = style.font
    font.name = 'Times New Roman'
    font.element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
    font.size = Pt(font_size_pt)
    font.bold = bold
    font.color.rgb = RGBColor(0, 0, 0) # Black

    # Paragraph settings
    pf = style.paragraph_format
    pf.alignment = align
    
    # Line spacing
    if line_spacing == 1.5:
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line_spacing
        
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    
    # Indentation
    if indent_chars > 0:
        # Approximate 1 char = font size. 
        # For mixed fonts, word counts char units differently, but usually based on body font (12pt).
        # DMU template says "2 characters". 2 * 12pt = 24pt.
        pf.first_line_indent = Pt(indent_chars * 12) 
    else:
        pf.first_line_indent = Pt(0)

def create_dmu_reference_doc(filename):
    doc = Document()
    
    # --- Page Setup (A4, Margins) ---
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(3.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    # Header/Footer distance
    section.header_distance = Cm(2.5)
    section.footer_distance = Cm(2.0)

    # --- 1. Normal (正文) ---
    # 宋体, 小四(12pt), 1.25倍行距, 首行缩进2字符
    update_style(doc, 'Normal', '宋体', 12, 
                 align=WD_ALIGN_PARAGRAPH.JUSTIFY, 
                 line_spacing=1.25, 
                 indent_chars=2)
    
    # --- 2. Headings ---
    
    # Heading 1 (Chapter Title - Numbered): 黑体, 小三(15pt), Left, 1.5 spacing, Before 0, After 1 line (~15pt)
    update_style(doc, 'Heading 1', '黑体', 15, 
                 align=WD_ALIGN_PARAGRAPH.LEFT, 
                 line_spacing=1.5, 
                 space_before_pt=0, 
                 space_after_pt=15, # 1 line at 15pt
                 bold=True,
                 heading_level=1)

    # Heading 2 (Section): 黑体, 四号(14pt), Left, 1.5 spacing, Before 0.5 line (~7pt), After 0
    update_style(doc, 'Heading 2', '黑体', 14, 
                 align=WD_ALIGN_PARAGRAPH.LEFT, 
                 line_spacing=1.5, 
                 space_before_pt=7, 
                 space_after_pt=0,
                 bold=True,
                 heading_level=2)

    # Heading 3 (Subsection): 黑体, 小四(12pt), Left, 1.5 spacing, Before 0.5 line (~6pt), After 0
    update_style(doc, 'Heading 3', '黑体', 12, 
                 align=WD_ALIGN_PARAGRAPH.LEFT, 
                 line_spacing=1.5, 
                 space_before_pt=6, 
                 space_after_pt=0,
                 bold=True,
                 heading_level=3)

    # --- 3. Custom Styles for Special Parts ---

    # Centered Heading (for Abstract, Intro, Conclusion, Refs)
    # Same as Heading 1 but Centered
    update_style(doc, 'CenteredHeading', '黑体', 15, 
                 align=WD_ALIGN_PARAGRAPH.CENTER, 
                 line_spacing=1.5, 
                 space_before_pt=0, 
                 space_after_pt=15,
                 bold=True)
                 
    # Abstract Title (Alias for CenteredHeading to match previous user code)
    update_style(doc, 'AbstractTitle', '黑体', 15, 
                 align=WD_ALIGN_PARAGRAPH.CENTER, 
                 line_spacing=1.5, 
                 space_before_pt=0, 
                 space_after_pt=15,
                 bold=True)

    # Abstract Body (Same as Normal but explicitly defined if needed, though Normal covers it)
    update_style(doc, 'AbstractBody', '宋体', 12, 
                 align=WD_ALIGN_PARAGRAPH.JUSTIFY, 
                 line_spacing=1.25, 
                 indent_chars=2)

    # Captions (Figure/Table)
    # 宋体, 五号(10.5pt), Center
    update_style(doc, 'Caption', '宋体', 10.5, 
                 align=WD_ALIGN_PARAGRAPH.CENTER, 
                 line_spacing=1.25,
                 bold=False)
                 
    # References Body
    # 宋体, 五号(10.5pt), 1.25 spacing, No indent (Left aligned)
    update_style(doc, 'ReferencesBody', '宋体', 10.5, 
                 align=WD_ALIGN_PARAGRAPH.LEFT, 
                 line_spacing=1.25, 
                 indent_chars=0)

    # Keywords Label (Character Style)
    # 黑体, 小四(12pt)
    try:
        style = doc.styles.add_style('KeywordsLabel', 2) # Character style
    except:
        style = doc.styles['KeywordsLabel']
    font = style.font
    font.name = 'Times New Roman'
    font.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    font.size = Pt(12)
    font.bold = True

    # --- 4. Add Sample Content ---
    doc.add_paragraph('Heading 1 (Chapter)', style='Heading 1')
    doc.add_paragraph('Normal text body. 宋体, 12pt, 1.25 spacing, 2 chars indent.', style='Normal')
    doc.add_paragraph('Heading 2 (Section)', style='Heading 2')
    doc.add_paragraph('Heading 3 (Subsection)', style='Heading 3')
    doc.add_paragraph('Centered Heading (Abstract)', style='CenteredHeading')
    doc.add_paragraph('Caption Text', style='Caption')
    doc.add_paragraph('Reference Item', style='ReferencesBody')

    doc.save(filename)
    print(f"Generated {filename} with DMU Thesis styles.")

if __name__ == '__main__':
    create_dmu_reference_doc('dmu_reference.docx')
